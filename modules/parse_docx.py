#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Parse a DOCX file to get dictionary words and output glossary table

Created on Mon Nov 05 2018

@author: dang

Written for Python 2.6
"""

from __future__ import unicode_literals
import docx
import sys
reload(sys)
sys.setdefaultencoding('utf8')
import csv
import re
import csv_tools

def check_series(text_list, set_list):
    """Check whether all items in each text_list[k] are in set_list
    """
    in_list = []
    for word in text_list:
        all_words = re.sub('\(.*?\)', ',', word).split(',')
        all_words = list(filter(None, all_words))
        component_in_list = [component.strip(' ') in set_list for component in all_words]
        this_word_in_list = all(component_in_list)
        in_list.append(this_word_in_list)
    return in_list


def check_field(text_list):
    with open('../ocr_tudien/word_fields.txt', 'rt') as listfile:
        set_list = listfile.read().split('\n')
        if set_list[-1] == '':
            del(set_list[-1])
    in_list = check_series(text_list, set_list)
    return in_list


def check_type(text_list):
    with open('../ocr_tudien/word_types.txt', 'rt') as listfile:
        set_list = listfile.read().split('\n')
        if set_list[-1] == '':
            del(set_list[-1])
    in_list = check_series(text_list, set_list)
    return in_list


def read_format(paragraph_style_bold, character_style_bolds, character_font_bolds, paragraph_style_italic, character_style_italics, character_font_italics):
    """Detect the character properties (bold, italic) based on info of the hierarchy:
       paragraph_style_bold (single value, from line.style.font.bold): on top level of the hierarchy, can be "None", "True", or "False"
       character_style_bolds (array, from docx line.runs[0].style.font.bold): second level, it inherits upper level if "None", overrides if "True" or "False"
       character_font_bolds (array, from docx line.runs[0].font.bold): third level, directly applied to character if "True" or "False", inherit if "None"
    """
    # Using style name doesn't work with different languages, example:
    # English Word file: Bold, Not Italic, German: Fett, Nicht kursiv
    if paragraph_style_bold == None:
        word_format_bolds = [False for item in character_font_bolds]
    else:
        word_format_bolds = [paragraph_style_bold for item in character_font_bolds]
    for k in range(len(character_style_bolds)):
        if character_style_bolds[k] != None:
            word_format_bolds[k] = character_style_bolds[k]
        if character_font_bolds[k] != None:
            word_format_bolds[k] = character_font_bolds[k]
    
    if paragraph_style_italic == None:
        word_format_italics = [False for item in character_font_italics]
    else:
        word_format_italics = [paragraph_style_italic for item in character_font_italics]
    for k in range(len(character_style_italics)):
        if character_style_italics[k] != None:
            word_format_italics[k] = character_style_italics[k]
        if character_font_italics[k] != None:
            word_format_italics[k] = character_font_italics[k]
    
    return word_format_bolds, word_format_italics


def remove_empty(list_text, *list_properties):
    properties = []
    for arg in list_properties: properties.append(arg)
    for k in range(len(list_text)-1, -1, -1):
        if list_text[k] == '':
            del(list_text[k])
            for prop in properties:
                del(prop[k])
    return tuple([list_text]+properties)


def join_unspaced(list_text, *list_properties):
    properties = []
    for arg in list_properties: properties.append(arg)
    for k_item in range(len(list_text)-1, 0, -1):
        if list_text[k_item][0] != ' ' and list_text[k_item-1][-1] != ' ':
            list_text[k_item-1] = ''.join([list_text[k_item-1], list_text[k_item]])
            del(list_text[k_item])
            for prop in properties:
                del(prop[k_item])
    return tuple([list_text]+properties)


def split_capital(phrase):
    """Split capital and small words.
    Capital words are only allowed at the first half of the phrase.
    If a capital word is after a small word, it is NOT marked capital.
    """
    format_capital = []
    #words = [phrase]
    #format_capital.append(words[0].isupper())
    
    #words = phrase.split()
    re_split_result = re.split('(\W)', phrase)
    words = list(filter(None, re_split_result))
    
    for k_item in range(len(words)-1, 0, -1):
            while len(words[k_item])>0 and words[k_item][0] == ',':
                words[k_item-1] += ','
                words[k_item] = words[k_item][1:]
    for k_item in range(len(words)-1, 0, -1):
        while len(words[k_item])>0 and words[k_item][0] == ' ':
            words[k_item-1] += ' '
            words[k_item] = words[k_item][1:]
        
    remove_empty(words)
        
    for k_item in range(len(words)-1, 0, -1):
        if words[k_item][0] != ' ' and words[k_item-1][-1] != ' ':
            words[k_item-1] = ''.join([words[k_item-1], words[k_item]])
            del(words[k_item])
    
    flag_prev_word_capital = True
    for k in range(len(words)):
        if words[k].isupper() and flag_prev_word_capital:
            if k>0 and (words[k-1].rstrip(' ')[-1] != ',' and words[k-1].rstrip(' ')[-1] != ')'):
                format_capital.append(False)
                flag_prev_word_capital = False
            else:
                format_capital.append(True)
            
        else:
            format_capital.append(False)
            flag_prev_word_capital = False
    return words, format_capital


def merge_similar_series(words, *list_properties):
    properties = []
    for arg in list_properties: properties.append(arg)
    for k in range(len(words)-1,0,-1):
        same_flag = True
        for type_format in properties:
            same_flag = same_flag and (type_format[k] == type_format[k-1])
        if same_flag:
            words[k-1] = ''.join([words[k-1], words[k]])
            del words[k]
            for type_format in properties:
                del type_format[k]
    return tuple([words]+properties)


def merge_with_comment_phrase(words, *list_properties):
    """If the starting word in phrase is in (), it should be a comment for the previous one,
    hence be added to the previous one.
    """
    properties = []
    for arg in list_properties: properties.append(arg)
    for k in range(len(words)-1,0,-1):
        if re.match('^(\(.*?\))', words[k].strip()):
            words[k-1] = ''.join([words[k-1], re.search('( *\(.*?\))', words[k].strip()).groups()[0]])
            words[k] = re.sub('(\(.*?\))', '', words[k])
            if words[k].strip(', ') == '':
                words[k-1] = ''.join([words[k-1], words[k]])
                del words[k]
                for type_format in properties:
                    del type_format[k]
    return tuple([words]+properties)


def re_parse(word_texts, word_format_bolds, word_format_italics):
    """Fix the leftover mistake in the data parsed, some example items:
    THAN gate road
    TOÁN, S_CHÊ figure 
    hình 
    (bản vẽ) - is a separate item, in fact it is a comment for previous item
    AB-Betrieb m Đ_TỬ dass AB mode chê'độ hạng AB, splitted into:
    ['AB-Betrieb', 'm', 'Đ_TỬ', 'dass AB', 'mode', 'chê'độ hạng', 'AB']
    """
    word_texts, word_format_bolds, word_format_italics = merge_with_comment_phrase(word_texts, word_format_bolds, word_format_italics)
    word_texts, word_format_bolds, word_format_italics = merge_similar_series(word_texts, word_format_bolds, word_format_italics)
    newword_texts = []
    newwordcase_capitals = []
    newword_format_bolds = []
    newword_format_italics = []
    # first word is assumed German, do not need to split capital
    if len(word_texts)>0:
        newword_texts.append(word_texts[0])
        newwordcase_capitals.append(False)
        newword_format_bolds.append(word_format_bolds[0])
        newword_format_italics.append(word_format_italics[0])
    for k in range(1, len(word_texts)):
        if word_format_italics[k]:
            words = [word_texts[k]]
            wordcase_capital = [False]
        else:
            words, wordcase_capital = split_capital(word_texts[k])
        newword_texts.extend(words)
        newwordcase_capitals.extend(wordcase_capital)
        for n in range(len(words)):
            newword_format_bolds.append(word_format_bolds[k])
            newword_format_italics.append(word_format_italics[k])
    
    newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = merge_with_comment_phrase(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
    newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = merge_similar_series(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
    return newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics

# encoding=utf8 by default
# See: https://stackoverflow.com/questions/21129020/how-to-fix-unicodedecodeerror-ascii-codec-cant-decode-byte

def read_docx(inputFile, exportFile = 'result.csv', logFile = 'log.txt'):

    extract_words = []
    extract_format_bolds = []
    extract_format_italics = []
    out_messages = []
    
    #number_items_in_line = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
    
    #number_blank_lines = 0
    #number_type_issues = 0
    #number_field_issues = 0
    #number_en_issues = 0
    #number_vi_issues = 0
    #number_total_item_issues = 0
    
    doc = docx.Document(inputFile)

    for k in range(len(doc.paragraphs)):
        line = doc.paragraphs[k]
        #print line.text
        # Make it compatible to Wordpad, this software removes all the styles
        paragraph_style_bold = None
        paragraph_style_italic = None
        if line.style != None:
            paragraph_style_bold = line.style.font.bold
            paragraph_style_italic = line.style.font.italic
        word_texts = [part.text for part in line.runs]
        character_style_bolds = [None for part in line.runs]
        character_style_italics = [None for part in line.runs]
        for k_item in range(len(line.runs)):
            part = line.runs[k_item]
            if part.style != None:
                character_style_bolds[k_item] = part.style.font.bold
                character_style_italics[k_item] = part.style.font.italic
        character_font_bolds = [part.font.bold for part in line.runs]
        character_font_italics = [part.font.italic for part in line.runs]
        
        word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics = remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == '&':
                word_texts[k_item-1] += '&'
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ',':
                word_texts[k_item-1] += ','
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ' ':
                word_texts[k_item-1] += ' '
                word_texts[k_item] = word_texts[k_item][1:]
        
        word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics = remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics = join_unspaced(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
            
        word_format_bolds, word_format_italics = read_format(paragraph_style_bold, character_style_bolds, character_font_bolds, paragraph_style_italic, character_style_italics, character_font_italics)
        #print(word_texts)
        #newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
        #print(newword_texts)
        
        extract_words.extend(word_texts)
        extract_format_bolds.extend(word_format_bolds)
        extract_format_italics.extend(word_format_italics)
        extract_words, extract_format_bolds, extract_format_italics = merge_similar_series(extract_words, extract_format_bolds, extract_format_italics)
        
        # Add end-of-line at the end of the paragraph, if the paragraph is non-empty
        if word_texts != []:
            extract_words.append('\n')
            extract_format_bolds.append(None)
            extract_format_italics.append(None)
        
    return extract_words, extract_format_bolds, extract_format_italics
    
    #excelFile = exportFile[:-4] + '.xls'
    #Csv_Excel.csv_to_xls(exportFile, excelFile)

def analyze_text(words, format_bolds, format_italics, de_condition={'bold':1, 'italic':-1}, en_condition={'bold':0, 'italic':-1}, end_of_item_eol=False, eoi_bold_to_unbold=False, eoi_unbold_to_bold=True, eoi_italic_to_unitalic=False, eoi_unitalic_to_italic=False):
    """
    Split the text into different categories, with the predifined condition for switching to new items
    The first group is named de_condition, second group is named en_condition, from historical purpose to parse Deutsch-English dictionaries
    """
    table = []
    de_words = ['']
    en_words = ['']
    if de_condition['bold'] != -1: de_condition['bold'] = not(not(de_condition['bold'])) # convert integer to boolean
    if de_condition['italic'] != -1: de_condition['italic'] = not(not(de_condition['italic']))
    
    last_format_bold = False
    last_format_italic = False
    de_bold_ok = False
    de_italic_ok = False
    en_bold_ok = False
    en_italic_ok = False
    item_no = 0
    for k in range(len(words)):
        if format_bolds[k] == False:
            new_format_bold = False
        elif format_bolds[k] == True:
            new_format_bold = True
        if format_italics[k] == False:
            new_format_italic = False
        elif format_italics[k] == True:
            new_format_italic = True

        # Check condition for switching to the new item
        flag_new = True
        if end_of_item_eol and words[k]!='\n': flag_new = False
        if eoi_bold_to_unbold and not(last_format_bold and not new_format_bold): flag_new = False
        if eoi_unbold_to_bold and not(not last_format_bold and new_format_bold): flag_new = False
        if eoi_italic_to_unitalic and not(last_format_italic and not new_format_italic): flag_new = False
        if eoi_unitalic_to_italic and not(not last_format_italic and new_format_italic): flag_new = False
        last_format_bold = new_format_bold
        last_format_italic = new_format_italic
        if flag_new and (de_words[item_no].strip() != '' or en_words[item_no].strip() != ''):
            de_words.append('')
            en_words.append('')
            item_no += 1
        
        if words[k]!='\n': # also format_bolds[k] and format_italics[k] must be either True or False
            if de_condition['bold'] != -1:
                de_bold_ok = not de_condition['bold']^format_bolds[k] # use XOR operator
            else: de_bold_ok = True
            if de_condition['italic'] != -1:
                de_italic_ok = not de_condition['italic']^format_bolds[k]
            else: de_italic_ok = True
            if de_bold_ok and de_italic_ok: de_words[item_no] += words[k]
        
            if en_condition['bold'] != -1:
                en_bold_ok = not en_condition['bold']^format_bolds[k]
            else: en_bold_ok = True
            if en_condition['italic'] != -1:
                en_italic_ok = not en_condition['italic']^format_bolds[k]
            else: en_italic_ok = True
            if en_bold_ok and en_italic_ok: en_words[item_no] += words[k]
        else:
            # Check if this end-of-line should be converted to a space, to connect lines
            if de_words[item_no] != '' and de_bold_ok and de_italic_ok: de_words[item_no] += ' '
            if en_words[item_no] != '' and en_bold_ok and en_italic_ok: en_words[item_no] += ' '            
        
    return de_words, en_words

def export_words_to_csv(de_words, en_words, exportFile = 'result.csv'):
    table = []
    for k in range(len(de_words)):
        table.append([de_words[k], en_words[k]])
    csv_tools.write_table_csv(exportFile, table)

def docx_to_csv(inputFile,  exportFile = 'result.csv', logFile = 'log.txt', **kwargs):
    """
    Example kwargs:
    group1_condition={'bold':1, 'italic':-1}, group2_condition={'bold':0, 'italic':-1}, eoi_eol=False, eoi_bold_to_unbold=False, eoi_unbold_to_bold=True, eoi_italic_to_unitalic=False, eoi_unitalic_to_italic=False
    """
    # Set default keys if not provided
    de_condition = kwargs['group1_condition'] if 'group1_condition' in kwargs else {'bold':1, 'italic':-1}
    en_condition = kwargs['group2_condition'] if 'group2_condition' in kwargs else {'bold':0, 'italic':-1}
    end_of_item_eol = kwargs['eoi_eol'] if 'eoi_eol' in kwargs else False
    eoi_bold_to_unbold = kwargs['eoi_bold_to_unbold'] if 'eoi_bold_to_unbold' in kwargs else False
    eoi_unbold_to_bold = kwargs['eoi_unbold_to_bold'] if 'eoi_unbold_to_bold' in kwargs else True
    eoi_italic_to_unitalic = kwargs['eoi_italic_to_unitalic'] if 'eoi_italic_to_unitalic' in kwargs else False
    eoi_unitalic_to_italic = kwargs['eoi_unitalic_to_italic'] if 'eoi_unitalic_to_italic' in kwargs else False

    extract_words, extract_format_bolds, extract_format_italics = read_docx(inputFile, exportFile, logFile)
    de_words, en_words = analyze_text(extract_words, extract_format_bolds, extract_format_italics, de_condition, en_condition, end_of_item_eol, eoi_bold_to_unbold, eoi_unbold_to_bold, eoi_italic_to_unitalic, eoi_unitalic_to_italic)
    export_words_to_csv(de_words, en_words, exportFile)
    # TODO: add logs


# Main operation, when calling: python Read_OCR.py input.docx output.csv
if __name__ == "__main__":
    inputFile = str(sys.argv[1])
    kwargs = {"group1_condition":{'bold':1, 'italic':-1}, "group2_condition":{'bold':0, 'italic':-1}, "eoi_eol":False, "eoi_bold_to_unbold":False, "eoi_unbold_to_bold":True, "eoi_italic_to_unitalic":False, "eoi_unitalic_to_italic":False}
    if len(sys.argv)>2:
        exportFile = str(sys.argv[2])
        docx_to_csv(inputFile, exportFile, group1_condition={'bold':1, 'italic':-1}, group2_condition={'bold':0, 'italic':-1}, eoi_eol=False, eoi_bold_to_unbold=False, eoi_unbold_to_bold=True, eoi_italic_to_unitalic=False, eoi_unitalic_to_italic=False)
    else:
        docx_to_csv(inputFile, **kwargs)
